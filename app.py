from flask import Flask, request, render_template, send_file
import numpy as np
import tensorflow as tf
import joblib
import os
import matplotlib
matplotlib.use('Agg')  # Use non-interactive backend
import matplotlib.pyplot as plt
import seaborn as sns
import pandas as pd
import warnings
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO

warnings.filterwarnings('ignore')

app = Flask(__name__)

# Load the model and the scaler
model = tf.keras.models.load_model('heart_disease_model.h5')
scaler = joblib.load('scaler.save')
data = pd.read_csv('ECG-Dataset.csv')

# Define column names in the order they appear in the CSV
FEATURE_NAMES = ['age', 'sex', 'smoke', 'years', 'ldl', 'chp', 'height', 'weight', 
                 'fh', 'active', 'lifestyle', 'ihd', 'hr', 'dm', 'bpsys', 'bpdias', 
                 'htn', 'ivsd', 'ecgpatt', 'qwave']

# Ensure the 'static' directory exists to save plots
if not os.path.exists('static'):
    os.makedirs('static')

def clean_old_plots():
    """Delete old plot files before generating new ones"""
    static_dir = 'static'
    if os.path.exists(static_dir):
        for filename in os.listdir(static_dir):
            if filename.endswith('.png'):
                try:
                    os.remove(os.path.join(static_dir, filename))
                except:
                    pass

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/predict', methods=['POST'])
def predict():
    try:
        # Clean old plots
        clean_old_plots()
        
        # Collect form data in order
        features = []
        for i in range(20):  # 20 features
            value = request.form.get(f'feature{i+1}')
            if value:
                features.append(float(value))
        
        if len(features) != 20:
            return "Error: Please fill all fields", 400
        
        final_features = np.array([features])
        
        # Standardize the input features
        final_features_scaled = scaler.transform(final_features)
        
        # Make prediction
        prediction_prob = model.predict(final_features_scaled, verbose=0)
        prediction_prob_value = float(prediction_prob[0][0])
        prediction = 1 if prediction_prob_value > 0.5 else 0
        confidence = prediction_prob_value if prediction == 1 else (1 - prediction_prob_value)
        
        # Generate prediction text
        if prediction == 1:
            prediction_text = f"⚠️ POSITIVE: Patient likely has heart disease (Confidence: {confidence*100:.1f}%)"
            risk_level = "HIGH RISK"
            color_code = "red"
        else:
            prediction_text = f"✓ NEGATIVE: Patient likely does not have heart disease (Confidence: {confidence*100:.1f}%)"
            risk_level = "LOW RISK"
            color_code = "green"
        
        # Create patient profile
        patient_profile = {
            'age': int(features[0]),
            'sex': 'Male' if int(features[1]) == 1 else 'Female',
            'smoke': 'Yes' if int(features[2]) == 1 else 'No',
            'ldl': round(features[4], 1),
            'hr': int(features[12]),
            'bpsys': int(features[14]),
            'bpdias': int(features[15]),
            'dm': 'Yes' if int(features[13]) == 1 else 'No',
            'htn': 'Yes' if int(features[16]) == 1 else 'No',
            'fh': 'Yes' if int(features[8]) == 1 else 'No',
            'bmi': round(features[7] / ((features[6]/100)**2), 1)  # weight / (height in m)^2
        }
        
        # Generate improved visualizations
        generate_visualizations(features, patient_profile)
        
        # Store prediction data in session/file for report generation
        prediction_data = {
            'prediction': prediction,
            'prediction_text': prediction_text,
            'risk_level': risk_level,
            'confidence': round(confidence*100, 1),
            'patient_profile': patient_profile,
            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }
        
        return render_template('result.html', 
                             prediction_text=prediction_text,
                             risk_level=risk_level,
                             color_code=color_code,
                             confidence=round(confidence*100, 1),
                             patient_profile=patient_profile,
                             prediction=prediction)
    
    except Exception as e:
        return f"Error: {str(e)}", 400

@app.route('/download-report', methods=['POST'])
def download_report():
    """Generate and download Word report"""
    try:
        # Get data from request
        data = request.json
        patient_profile = data.get('patient_profile', {})
        prediction_text = data.get('prediction_text', '')
        confidence = data.get('confidence', 0)
        risk_level = data.get('risk_level', '')
        
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Heart Disease Prediction Report', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add timestamp
        timestamp_para = doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%B %d, %Y at %H:%M:%S')}")
        timestamp_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        doc.add_paragraph()  # Blank line
        
        # Add prediction result
        doc.add_heading('Prediction Result', level=1)
        result_para = doc.add_paragraph(prediction_text)
        result_para.runs[0].font.size = Pt(14)
        result_para.runs[0].font.bold = True
        
        # Add confidence bar text
        doc.add_paragraph(f"Confidence Level: {confidence}%")
        doc.add_paragraph(f"Risk Assessment: {risk_level}")
        
        doc.add_paragraph()  # Blank line
        
        # Add patient profile
        doc.add_heading('Patient Health Profile', level=1)
        
        table = doc.add_table(rows=11, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Set header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Parameter'
        header_cells[1].text = 'Value'
        
        # Fill data
        profile_data = [
            ('Age', f"{patient_profile.get('age', 'N/A')} years"),
            ('Sex', patient_profile.get('sex', 'N/A')),
            ('LDL Cholesterol', f"{patient_profile.get('ldl', 'N/A')} mg/dL"),
            ('Heart Rate', f"{patient_profile.get('hr', 'N/A')} bpm"),
            ('Blood Pressure', f"{patient_profile.get('bpsys', 'N/A')}/{patient_profile.get('bpdias', 'N/A')} mmHg"),
            ('BMI', patient_profile.get('bmi', 'N/A')),
            ('Smoker', patient_profile.get('smoke', 'N/A')),
            ('Diabetes', patient_profile.get('dm', 'N/A')),
            ('Hypertension', patient_profile.get('htn', 'N/A')),
            ('Family History', patient_profile.get('fh', 'N/A'))
        ]
        
        for idx, (param, value) in enumerate(profile_data, start=1):
            row_cells = table.rows[idx].cells
            row_cells[0].text = param
            row_cells[1].text = str(value)
        
        doc.add_paragraph()  # Blank line
        
        # Add risk assessment
        doc.add_heading('Risk Factor Summary', level=1)
        risk_factors = []
        if patient_profile.get('dm') == 'Yes':
            risk_factors.append('• Diabetes')
        if patient_profile.get('htn') == 'Yes':
            risk_factors.append('• Hypertension')
        if patient_profile.get('smoke') == 'Yes':
            risk_factors.append('• Smoking')
        if patient_profile.get('fh') == 'Yes':
            risk_factors.append('• Family History')
        
        if risk_factors:
            for factor in risk_factors:
                doc.add_paragraph(factor)
        else:
            doc.add_paragraph("No major risk factors identified")
        
        doc.add_paragraph()  # Blank line
        
        # Add visualizations section
        doc.add_heading('Analysis Visualizations', level=1)
        
        if os.path.exists('static/patient_metrics.png'):
            doc.add_heading('Patient Metrics Dashboard', level=2)
            doc.add_picture('static/patient_metrics.png', width=Inches(6))
        
        if os.path.exists('static/correlation_heatmap.png'):
            doc.add_heading('Feature Correlations', level=2)
            doc.add_picture('static/correlation_heatmap.png', width=Inches(6))
        
        if os.path.exists('static/risk_distribution.png'):
            doc.add_heading('Risk Distribution in Dataset', level=2)
            doc.add_picture('static/risk_distribution.png', width=Inches(6))
        
        doc.add_paragraph()  # Blank line
        
        # Add footer with disclaimer
        doc.add_heading('Disclaimer', level=2)
        disclaimer = doc.add_paragraph(
            "This report is generated by an AI-based heart disease prediction model. "
            "It should be used as a supplementary tool and not as a substitute for professional medical advice. "
            "Always consult with a qualified healthcare professional for accurate diagnosis and treatment."
        )
        disclaimer.runs[0].font.size = Pt(10)
        disclaimer.runs[0].italic = True
        
        # Save to bytes
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        # Return as downloadable file
        filename = f"Heart_Disease_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        return send_file(
            doc_bytes,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
    
    except Exception as e:
        return f"Error generating report: {str(e)}", 400

def generate_visualizations(features, patient_profile):
    """Generate clear, concise visualizations"""
    
    sns.set_style("darkgrid")
    plt.rcParams['figure.facecolor'] = '#f5f5f5'
    
    # 1. Patient Key Metrics Dashboard
    fig, axes = plt.subplots(2, 3, figsize=(14, 8))
    fig.suptitle('Patient Health Metrics Summary', fontsize=16, fontweight='bold')
    
    # Age
    ax = axes[0, 0]
    ax.barh(['Age'], [patient_profile['age']], color='#3498db', edgecolor='black', linewidth=2)
    ax.set_xlim(0, 100)
    ax.text(patient_profile['age'] + 2, 0, f"{patient_profile['age']} yrs", va='center', fontweight='bold')
    ax.set_xlabel('Years')
    
    # LDL Cholesterol
    ax = axes[0, 1]
    color = '#e74c3c' if patient_profile['ldl'] > 130 else '#f39c12' if patient_profile['ldl'] > 100 else '#2ecc71'
    ax.barh(['LDL'], [patient_profile['ldl']], color=color, edgecolor='black', linewidth=2)
    ax.set_xlim(0, 300)
    ax.text(patient_profile['ldl'] + 5, 0, f"{patient_profile['ldl']} mg/dL", va='center', fontweight='bold')
    ax.set_xlabel('mg/dL')
    
    # Heart Rate
    ax = axes[0, 2]
    color = '#e74c3c' if patient_profile['hr'] > 100 else '#f39c12' if patient_profile['hr'] > 85 else '#2ecc71'
    ax.barh(['HR'], [patient_profile['hr']], color=color, edgecolor='black', linewidth=2)
    ax.set_xlim(0, 160)
    ax.text(patient_profile['hr'] + 2, 0, f"{patient_profile['hr']} bpm", va='center', fontweight='bold')
    ax.set_xlabel('Beats/min')
    
    # Systolic BP
    ax = axes[1, 0]
    color = '#e74c3c' if patient_profile['bpsys'] > 140 else '#f39c12' if patient_profile['bpsys'] > 120 else '#2ecc71'
    ax.barh(['Systolic BP'], [patient_profile['bpsys']], color=color, edgecolor='black', linewidth=2)
    ax.set_xlim(0, 220)
    ax.text(patient_profile['bpsys'] + 3, 0, f"{patient_profile['bpsys']} mmHg", va='center', fontweight='bold')
    ax.set_xlabel('mmHg')
    
    # BMI
    ax = axes[1, 1]
    color = '#e74c3c' if patient_profile['bmi'] > 30 else '#f39c12' if patient_profile['bmi'] > 25 else '#2ecc71'
    ax.barh(['BMI'], [patient_profile['bmi']], color=color, edgecolor='black', linewidth=2)
    ax.set_xlim(0, 40)
    ax.text(patient_profile['bmi'] + 0.5, 0, f"{patient_profile['bmi']}", va='center', fontweight='bold')
    ax.set_xlabel('kg/m²')
    
    # Risk Factors Summary
    ax = axes[1, 2]
    ax.axis('off')
    risk_factors = []
    if patient_profile['dm'] == 'Yes':
        risk_factors.append('• Diabetes')
    if patient_profile['htn'] == 'Yes':
        risk_factors.append('• Hypertension')
    if patient_profile['smoke'] == 'Yes':
        risk_factors.append('• Smoking')
    if patient_profile['fh'] == 'Yes':
        risk_factors.append('• Family History')
    
    if not risk_factors:
        risk_text = "No major risk factors"
        ax.text(0.5, 0.5, risk_text, ha='center', va='center', fontsize=12, 
               bbox=dict(boxstyle='round', facecolor='#2ecc71', alpha=0.7), fontweight='bold')
    else:
        risk_text = "Risk Factors:\n" + "\n".join(risk_factors)
        ax.text(0.5, 0.5, risk_text, ha='center', va='center', fontsize=11, 
               bbox=dict(boxstyle='round', facecolor='#e74c3c', alpha=0.7), fontweight='bold')
    
    plt.tight_layout()
    plt.savefig('static/patient_metrics.png', dpi=100, bbox_inches='tight')
    plt.close()
    
    # 2. Correlation Heatmap (Dataset overview - simplified)
    fig, ax = plt.subplots(figsize=(10, 8))
    
    # Select key features for correlation
    key_features = ['age', 'ldl', 'hr', 'bpsys', 'bpdias', 'weight', 'target']
    corr_matrix = data[key_features].corr()
    
    sns.heatmap(corr_matrix, annot=True, fmt='.2f', cmap='RdYlGn', center=0, 
                square=True, linewidths=1, cbar_kws={"shrink": 0.8}, ax=ax)
    ax.set_title('Feature Correlations with Heart Disease Prediction', fontsize=14, fontweight='bold')
    
    plt.tight_layout()
    plt.savefig('static/correlation_heatmap.png', dpi=100, bbox_inches='tight')
    plt.close()
    
    # 3. Risk Distribution Comparison
    fig, ax = plt.subplots(figsize=(10, 6))
    
    risk_labels = ['No Heart Disease', 'Has Heart Disease']
    risk_counts = data['target'].value_counts().sort_index()
    colors = ['#2ecc71', '#e74c3c']
    
    bars = ax.bar(risk_labels, risk_counts.values, color=colors, edgecolor='black', linewidth=2, width=0.6)
    ax.set_ylabel('Number of Patients', fontsize=12, fontweight='bold')
    ax.set_title('Heart Disease Distribution in Dataset', fontsize=14, fontweight='bold')
    
    # Add value labels on bars
    for bar in bars:
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2., height,
               f'{int(height)}\n({int(height)/len(data)*100:.1f}%)',
               ha='center', va='bottom', fontweight='bold')
    
    plt.tight_layout()
    plt.savefig('static/risk_distribution.png', dpi=100, bbox_inches='tight')
    plt.close()

if __name__ == "__main__":
    app.run(debug=False, host='127.0.0.1', port=5000)
